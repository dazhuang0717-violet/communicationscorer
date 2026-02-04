import streamlit as st
import pandas as pd
import numpy as np
import requests
from bs4 import BeautifulSoup
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
import io
import math
import json
import re
import time

st.set_page_config(
    page_title="肿瘤业务-传播价值 AI 评分系统",
    page_icon="📡",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
    <style>
        [data-testid="stAppViewContainer"] { background-color: #ffffff !important; color: #31333F !important; }
        [data-testid="stSidebar"] { background-color: #f8f9fa !important; border-right: 1px solid #e0e0e0; }
        
        header[data-testid="stHeader"] { background-color: #ffffff !important; border-bottom: 1px solid #f0f2f6; }
        header[data-testid="stHeader"] button, header[data-testid="stHeader"] a, header[data-testid="stHeader"] svg { color: #31333F !important; fill: #31333F !important; }
        
        .stTextInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] { 
            color: #31333F !important; 
            background-color: #ffffff !important; 
            border: 1px solid #d1d5db; 
        }
        
        .stTextInput input:focus, .stTextArea textarea:focus { 
            border-color: #1E88E5 !important; 
            box-shadow: 0 0 0 1px #1E88E5 !important;
        }
        div[data-baseweb="select"] > div:focus-within {
            border-color: #1E88E5 !important;
        }

        button[kind="primary"] {
            background-color: #1E88E5 !important;
            border-color: #1E88E5 !important;
        }
        button[kind="secondary"] {
            border-color: #1E88E5 !important;
            color: #1E88E5 !important;
        }
        
        [data-testid="stFileUploaderDropzone"] { background-color: #f8f9fa !important; border: 1px dashed #d1d5db !important; }
        [data-testid="stFileUploaderDropzone"] div, [data-testid="stFileUploaderDropzone"] span, [data-testid="stFileUploaderDropzone"] p { color: #31333F !important; }
        
        [data-testid="stDataFrame"] { 
            color: #000000 !important; 
        }
        [data-testid="stDataFrame"] svg { fill: #31333F !important; }
        
        [data-testid="stDataFrame"] * {
            font-family: "Microsoft YaHei", "PingFang SC", "Source Sans Pro", sans-serif !important;
        }
        
        #MainMenu { visibility: hidden; }
        footer { visibility: hidden; }
        
        .stAlert { 
            background-color: #e3f2fd !important; 
            border: 1px solid #90caf9 !important; 
            color: #0d47a1 !important; 
        }

        .stTabs [data-baseweb="tab-list"] button[aria-selected="true"] {
            border-bottom-color: #1E88E5 !important;
        }
        .stTabs [data-baseweb="tab-list"] button[aria-selected="true"] p {
            color: #1E88E5 !important;
        }
        div[data-baseweb="tab-highlight"] {
            background-color: #1E88E5 !important;
        }
        
        .stTabs [data-baseweb="tab-list"] button:hover p {
            color: #1E88E5 !important;
        }
        .stTabs [data-baseweb="tab-list"] button:hover {
            color: #1E88E5 !important;
            border-bottom-color: #1E88E5 !important;
        }
    </style>
""", unsafe_allow_html=True)

class ScorerEngine:
    def __init__(self, key):
        self.api_key = key
        self.portkey_url = "https://api.portkey.ai/v1/chat/completions"

    def read_docx_content(self, file_obj):
        try:
            file_obj.seek(0)
            doc = Document(file_obj)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip(): full_text.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip(): full_text.append(para.text.strip())
            return "\n".join(full_text)
        except Exception as e:
            return f"Error: {str(e)}"

    def fetch_url_content(self, url):
        if not url or pd.isna(url): return ""
        if not str(url).startswith('http'): return ""
        try:
            jina_url = f"https://r.jina.ai/{url}"
            response = requests.get(jina_url, timeout=5)
            if response.status_code == 200 and len(response.text) > 50: return response.text[:10000]
        except: pass 
        try:
            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(url, headers=headers, timeout=5)
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                text = " ".join([p.get_text() for p in soup.find_all('p')])
                if len(text) > 50: return text[:10000]
        except: pass
        return ""

    def calculate_volume_quality(self, views, interactions):
        try:
            def clean_num(x):
                if isinstance(x, str):
                    x = re.sub(r'[kK]', '000', x)
                    x = re.sub(r'[^\d\.]', '', x)
                return float(x) if x else 0.0
            v = clean_num(views)
            i = clean_num(interactions)
            raw_score = math.log10(v + i * 5 + 1) * 1.5
            return min(10.0, round(raw_score, 1))
        except: return 0.0

    def get_media_tier_score(self, media_name, tiers_config):
        if not media_name or pd.isna(media_name): return 3
        m_name = str(media_name).lower().strip()
        for tier_name, tier_list in tiers_config.items():
            for configured_media in tier_list:
                if configured_media and configured_media in m_name:
                    if tier_name == 'tier1': return 10
                    if tier_name == 'tier2': return 8
                    if tier_name == 'tier3': return 5
        return 3

    def analyze_content_with_ai(self, content, key_message, project_desc, audience_mode, media_name):
        if not self.api_key or not str(self.api_key).strip(): return 0, 0, 0, "API Key Missing", "无评价"
        
        if not content or len(str(content).strip()) < 10:
             return 0, 0, 0, "内容过短/无效", "内容过短，无法生成评价"

        safe_km = key_message if key_message else "文章主题及核心观点"
        safe_desc = project_desc if project_desc else "一般性行业项目"

        prompt = f"""
        你是一个专业的公关传播分析师。请严格按照以下规则对内容进行评分：

        【评分规则】
        1. **信息匹配 (km_score)**: 请仔细阅读【待分析文本】，判断其是否有效传递了【核心传播信息】。
        2. **获客效能 (acquisition_score)**: 基于【项目描述】，评估这个项目的获客效能。
        3. **受众精准度 (audience_precision_score)**: 仅根据【媒体名称】和【目标受众模式】进行判断。例如，如果是"HCP"模式但媒体是大众娱乐媒体，则分数应较低。

        【输入信息】
        - 目标受众模式: {audience_mode}
        - 媒体名称: {media_name}
        - 核心传播信息 (Key Message): {safe_km}
        - 项目描述: {safe_desc}
        - 待分析文本: 
        {content[:3000]}... (内容截断)

        【输出任务】
        请返回 JSON 格式的分数（0-10分）以及一段简短评价，格式如下：
        {{
            "km_score": <分数>,
            "acquisition_score": <分数>,
            "audience_precision_score": <分数>,
            "comment": "简短评价：客观指出优缺点，概括性强，100字以内。"
        }}
        """
        
        candidate_models = [
            'gemini-2.0-flash', 
            'gemini-1.5-flash-latest',
            'gemini-1.5-pro-latest',
            'gemini-1.5-flash',
            'gemini-1.5-pro'
        ]
        
        headers = {
            "x-portkey-api-key": self.api_key,
            "x-portkey-provider": "google",
            "Content-Type": "application/json"
        }

        def extract_json(text):
            try: return json.loads(text)
            except: pass
            try:
                clean = text.replace('```json', '').replace('```', '').strip()
                return json.loads(clean)
            except: pass
            try:
                match = re.search(r'\{.*\}', text, re.DOTALL)
                if match: return json.loads(match.group(0))
            except: pass
            return None

        last_error = None
        for model_name in candidate_models:
            try:
                payload = {
                    "model": model_name,
                    "messages": [{"role": "user", "content": prompt}]
                }
                response = requests.post(self.portkey_url, headers=headers, json=payload, timeout=30)
                
                if response.status_code == 200:
                    res_json = response.json()
                    res_text = res_json['choices'][0]['message']['content']
                    data = extract_json(res_text)
                    if data:
                        return (
                            data.get('km_score', 0), 
                            data.get('acquisition_score', 0), 
                            data.get('audience_precision_score', 0), 
                            "Success",
                            data.get('comment', 'AI 未返回评价')
                        )
                elif response.status_code == 412:
                    last_error = f"Model {model_name} not allowed"
                    continue
                else:
                    raise ValueError(f"HTTP {response.status_code}: {response.text}")
            except Exception as e:
                last_error = e
                if "429" in str(e): 
                    time.sleep(1)
                    continue
                continue

        return 0, 0, 0, f"AI Failed ({str(last_error)})", "AI 调用失败，请检查 Portkey 权限"

def generate_html_report(project_name, metrics, charts, df_top):
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{project_name} - 评分报告</title>
        <style>
            body {{ font-family: "Microsoft YaHei", sans-serif; padding: 40px; color: #333; }}
            h1 {{ color: #1E88E5; border-bottom: 2px solid #1E88E5; padding-bottom: 10px; }}
            h2 {{ color: #1E88E5; margin-top: 30px; }}
            .metrics-container {{ display: flex; justify-content: space-between; margin-bottom: 30px; background: #f8f9fa; padding: 20px; border-radius: 8px; }}
            .metric-box {{ text-align: center; }}
            .metric-val {{ font-size: 24px; font-weight: bold; color: #1E88E5; }}
            .metric-lbl {{ font-size: 14px; color: #666; }}
            .chart-container {{ margin-bottom: 40px; page-break-inside: avoid; }}
            table {{ width: 100%; border-collapse: collapse; margin-top: 20px; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
            th {{ background-color: #1E88E5; color: white; }}
            tr:nth-child(even) {{ background-color: #f2f2f2; }}
            @media print {{
                .no-print {{ display: none; }}
                body {{ padding: 0; }}
            }}
        </style>
    </head>
    <body>
        <h1>📈 项目评分报告: {project_name}</h1>
        
        <div class="metrics-container">
            <div class="metric-box"><div class="metric-val">{metrics['total']:.2f}</div><div class="metric-lbl">项目总分</div></div>
            <div class="metric-box"><div class="metric-val">{metrics['demand']:.2f}</div><div class="metric-lbl">真需求</div></div>
            <div class="metric-box"><div class="metric-val">{metrics['acquisition']:.2f}</div><div class="metric-lbl">获客效能</div></div>
            <div class="metric-box"><div class="metric-val">{metrics['volume']:.2f}</div><div class="metric-lbl">声量</div></div>
        </div>

        <h2>📊 数据洞察</h2>
        <div style="display: flex; flex-wrap: wrap;">
            <div style="width: 50%; min-width: 300px;" class="chart-container">
                <h3>项目能力雷达</h3>
                {charts['radar']}
            </div>
            <div style="width: 50%; min-width: 300px;" class="chart-container">
                <h3>传播价值矩阵</h3>
                {charts['scatter']}
            </div>
        </div>
        <div class="chart-container">
            <h3>媒体贡献 TOP 榜单</h3>
            {charts['bar']}
        </div>

        <h2>🏆 详细数据 (Top 10)</h2>
        {df_top.to_html(index=False)}

        <div class="no-print" style="margin-top: 40px; text-align: center; color: #888;">
            <p>💡 提示: 请使用浏览器菜单 "文件" -> "打印" (或 Ctrl+P)，选择 "另存为 PDF"。</p>
        </div>
    </body>
    </html>
    """
    return html_content

with st.sidebar:
    st.header("⚙️ 系统配置")
    
    api_key = st.text_input("🔑 Portkey API Key", value="", type="password")

    st.subheader("📋 项目基础信息")
    project_name = st.text_input("项目名称")
    project_key_message = st.text_input("核心信息 (Key Message)", value="")
    project_desc = st.text_area("项目描述 (用于评估获客)", value="", height=100)
    audience_mode = st.radio("目标受众模式", ["大众 (General)", "患者 (Patient)", "医疗专业人士 (HCP)"])

    st.markdown("---")
    st.subheader("🏆 媒体分级")
    st.caption("输入媒体名称，用逗号分隔")
    tier1_input = st.text_area("Tier 1 (10分)", value="", height=68)
    tier2_input = st.text_area("Tier 2 (8分)", value="", height=68)
    tier3_input = st.text_area("Tier 3 (5分)", value="", height=68)

    def parse_tiers(text):
        return [x.strip().lower() for x in text.split(',') if x.strip()]
    
    tier_config = {
        'tier1': parse_tiers(tier1_input),
        'tier2': parse_tiers(tier2_input),
        'tier3': parse_tiers(tier3_input)
    }

engine = ScorerEngine(api_key)

st.title("📡 肿瘤业务-传播价值 AI 评分系统")

with st.expander("查看核心算法公式", expanded=False):
    st.markdown("""
    <div style="text-align: center; font-size: 20px; line-height: 2.5; color: #31333F; background-color: #f8f9fa; padding: 20px; border-radius: 10px; font-family: sans-serif;">
        <span style="font-weight: bold; color: #1E88E5;">总分</span> = 0.5 × 真需求 + 0.2 × 获客效能 + 0.3 × 声量<br>
        <span style="font-weight: bold; color: #1E88E5;">真需求</span> = 0.6 × 信息匹配 + 0.4 × 受众精准度 &nbsp;&nbsp;&nbsp;&nbsp;&nbsp; <span style="font-weight: bold; color: #1E88E5;">声量</span> = 0.6 × 传播质量 + 0.4 × 媒体分级
    </div>
    """, unsafe_allow_html=True)

tab1, tab2, tab3 = st.tabs(["📄 新闻稿评分", "📊 媒体报道评分", "📈 项目评分"])

with tab1:
    st.info("📄 上传新闻稿 Word 文档，AI 将评价核心信息传递情况。")
    uploaded_word = st.file_uploader("上传 .docx 文件", type=['docx'])
    
    if 'word_analysis_result' not in st.session_state:
        st.session_state.word_analysis_result = None

    if uploaded_word:
        st.info("✅ 文档已就绪")
        
        if st.button("开始分析", key="btn_word_analyze"):
            if not api_key:
                st.error("❌ 请先在侧边栏输入 Portkey API Key")
            elif not project_key_message:
                st.warning("⚠️ 请在左侧填写【核心信息】")
            else:
                with st.spinner("AI 正在阅读文档..."):
                    try:
                        full_text = engine.read_docx_content(uploaded_word)
                        if len(full_text.strip()) < 10:
                            st.error(f"文档内容过少 (提取到 {len(full_text)} 字)，无法进行分析。")
                            st.session_state.word_analysis_result = None
                        else:
                            km, acq, prec, status, comment = engine.analyze_content_with_ai(
                                full_text, project_key_message, project_desc, audience_mode, "内部稿件"
                            )
                            st.session_state.word_analysis_result = {
                                "km": km, 
                                "status": status, 
                                "text_len": len(full_text),
                                "comment": comment
                            }
                    except Exception as e:
                        st.error(f"解析错误: {e}")
    
    if st.session_state.word_analysis_result:
        res = st.session_state.word_analysis_result
        st.divider()
        if res['km'] > 0:
            st.metric("核心信息匹配度", f"{res['km']}/10")
            st.progress(res['km']/10)
            
            st.success("分析成功！")
            
            st.markdown(f"""
            <div style="background-color: #f0f8ff; padding: 15px; border-radius: 8px; border-left: 5px solid #1E88E5;">
                <h4 style="color: #1E88E5; margin-top: 0;">💡 AI 简评</h4>
                <p style="color: #31333F;">{res.get('comment', '暂无评价')}</p>
            </div>
            """, unsafe_allow_html=True)
            
        else:
            st.error(f"评分失败 (0分)。\n原因: {res['status']}")

if 'batch_results_df' not in st.session_state:
    st.session_state.batch_results_df = None

with tab2:
    st.info("💡 微信公众号、视频号等封闭平台内容无法自动爬取，请在 Excel 中插入“正文”列并手动填入文章内容。")
    
    uploaded_file = st.file_uploader("上传媒体监测报表", type=['xlsx', 'csv'])

    if uploaded_file:
        try:
            if uploaded_file.name.endswith('.csv'):
                try: df = pd.read_csv(uploaded_file)
                except: uploaded_file.seek(0); df = pd.read_csv(uploaded_file, encoding='gbk')
            else:
                df = pd.read_excel(uploaded_file)
            
            df.columns = df.columns.str.strip()

            if '媒体' in df.columns and '媒体名称' not in df.columns: df['媒体名称'] = df['媒体']
            if '链接' in df.columns and 'URL' not in df.columns: df['URL'] = df['链接']

            def to_num(x):
                try:
                    if pd.isna(x) or x == '': return 0.0
                    s = str(x).replace(',', '').replace('+', '').strip()
                    if '万' in s: return float(s.replace('万', '')) * 10000
                    return float(s)
                except: return 0.0

            if 'PV' not in df.columns: df['PV'] = 0
            if '浏览量' not in df.columns: df['浏览量'] = 0

            df['Clean_Views'] = df['PV'].apply(to_num)
            mask = df['Clean_Views'] == 0
            df.loc[mask, 'Clean_Views'] = df.loc[mask, '浏览量'].apply(to_num)
            df['浏览量'] = df['Clean_Views']

            df['互动量'] = 0
            for col in ['点赞量', '评论量', '转发量']:
                if col in df.columns: df['互动量'] += df[col].apply(to_num)

            required_cols = ['媒体名称', 'URL', '互动量', '浏览量']
            missing_cols = [col for col in required_cols if col not in df.columns]
            
            if missing_cols:
                st.error(f"⚠️ 文件缺少必要列: {missing_cols}")
            else:
                df.index = range(1, len(df) + 1)
                st.info(f"✅ 成功读取 {len(df)} 条数据，以下为预览:")
                
                preview_cols_candidates = ['标题', '媒体', '媒体类型', '浏览量', '互动量', '链接']
                actual_preview_cols = [c for c in preview_cols_candidates if c in df.columns]
                
                if actual_preview_cols:
                    preview_df = df[actual_preview_cols].copy()
                    st.dataframe(preview_df, use_container_width=True)
                else:
                    st.dataframe(df, use_container_width=True)
                
                st.markdown("---")
                
                if st.button("开始分析", key="btn_xlsx_analyze"):
                    if not api_key:
                        st.error("❌ 请先在侧边栏配置 Portkey API Key")
                    else:
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        results = []
                        total_rows = len(df)

                        for index, row in df.iterrows():
                            status_text.text(f"⏳ 正在分析第 {index}/{total_rows} 条: {row['媒体名称']}...")
                            
                            vol_quality = engine.calculate_volume_quality(row['浏览量'], row['互动量'])
                            tier_score = engine.get_media_tier_score(row['媒体名称'], tier_config)
                            volume_total = 0.6 * vol_quality + 0.4 * tier_score
                            
                            content = ""
                            if '正文' in df.columns and pd.notna(row['正文']):
                                content = str(row['正文'])
                                msg_suffix = " (基于Excel文本)"
                            elif 'Content' in df.columns and pd.notna(row['Content']):
                                content = str(row['Content'])
                                msg_suffix = " (基于Excel文本)"
                            else:
                                content = engine.fetch_url_content(row['URL'])
                                msg_suffix = ""

                            if not content and '标题' in df.columns and pd.notna(row['标题']):
                                content = f"文章标题：{row['标题']}"
                                msg_suffix = " (基于标题)"
                            
                            if content:
                                km_score, acq_score, prec_score, msg, _ = engine.analyze_content_with_ai(
                                    content, project_key_message, project_desc, audience_mode, row['媒体名称']
                                )
                                msg += msg_suffix
                            else:
                                km_score, acq_score, prec_score = 0, 0, 0
                                msg = "无内容"

                            true_demand = 0.6 * km_score + 0.4 * prec_score
                            total_score = (0.5 * true_demand) + (0.2 * acq_score) + (0.3 * volume_total)

                            results.append({
                                "媒体名称": row['媒体名称'],
                                "项目总分": round(total_score, 2),
                                "真需求": round(true_demand, 2),
                                "获客效能": acq_score,
                                "声量": round(volume_total, 2),
                                "声量小分": round(volume_total, 2),
                                "核心信息匹配": km_score,
                                "受众精准度": prec_score, 
                                "媒体分级": tier_score,
                                "传播质量": vol_quality,
                                "状态": msg
                            })
                            progress_bar.progress(index / total_rows)

                        status_text.info("🎉 分析完成！")
                        
                        res_df = pd.DataFrame(results)
                        res_df.index = range(1, len(res_df) + 1)
                        st.session_state.batch_results_df = res_df
        
        except Exception as e:
            st.error(f"文件处理错误: {e}")

    if st.session_state.batch_results_df is not None:
        res_df = st.session_state.batch_results_df
        st.divider()
        st.subheader("📋 媒体报道评分")
        tab2_cols = ['媒体名称', '媒体分级', '受众精准度', '传播质量', '声量']
        st.dataframe(res_df[tab2_cols], use_container_width=True)
        
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            res_df.to_excel(writer, index=True)
        
        st.download_button(
            label="📥 导出评分报告 (Excel)",
            data=buffer.getvalue(),
            file_name=f"{project_name}_scoring_report.xlsx" if project_name else "scoring_report.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            type="primary"
        )

with tab3:
    if st.session_state.batch_results_df is None:
        st.info("👋 请先完成“新闻稿评分”和“媒体报道评分”。")
    else:
        res_df = st.session_state.batch_results_df
        
        st.subheader(f"📈 项目评分: {project_name if project_name else '未命名项目'}")
        
        m1, m2, m3, m4 = st.columns(4)
        avg_score = res_df['项目总分'].mean()
        metrics = {
            'total': avg_score,
            'demand': res_df['真需求'].mean(),
            'acquisition': res_df['获客效能'].mean(),
            'volume': res_df['声量'].mean()
        }
        
        m1.metric("项目总分", f"{metrics['total']:.2f}")
        m2.metric("真需求", f"{metrics['demand']:.2f}")
        m3.metric("获客效能", f"{metrics['acquisition']:.2f}")
        m4.metric("声量", f"{metrics['volume']:.2f}")

        st.divider()
        st.subheader("📊 数据洞察")

        col_chart1, col_chart2 = st.columns(2)
        
        charts = {}

        with col_chart1:
            st.markdown("##### 🕸️ 项目雷达")
            radar_categories = ['核心信息匹配', '获客效能', '受众精准度', '媒体分级', '传播质量']
            radar_values = [
                res_df['核心信息匹配'].mean(),
                res_df['获客效能'].mean(),
                res_df['受众精准度'].mean(),
                res_df['媒体分级'].mean(),
                res_df['传播质量'].mean()
            ]
            
            fig_radar = go.Figure()
            fig_radar.add_trace(go.Scatterpolar(
                r=radar_values,
                theta=radar_categories,
                fill='toself',
                name='项目平均表现',
                line_color='#1E88E5',
                fillcolor='rgba(30, 136, 229, 0.3)'
            ))
            fig_radar.update_layout(
                polar=dict(
                    radialaxis=dict(visible=True, range=[0, 10])
                ),
                showlegend=False,
                margin=dict(l=40, r=40, t=30, b=30),
                height=350
            )
            st.plotly_chart(fig_radar, use_container_width=True)
            charts['radar'] = fig_radar.to_html(full_html=False, include_plotlyjs='cdn')

        with col_chart2:
            st.markdown("##### 💠 传播矩阵")
            fig_scatter = px.scatter(
                res_df,
                x='声量',
                y='真需求',
                color='项目总分',
                hover_data=['媒体名称'],
                size='项目总分', 
                color_continuous_scale='Blues',
                height=350
            )
            fig_scatter.update_layout(margin=dict(l=20, r=20, t=30, b=20))
            st.plotly_chart(fig_scatter, use_container_width=True)
            charts['scatter'] = fig_scatter.to_html(full_html=False, include_plotlyjs='cdn')

        st.markdown("##### 🏆 媒体榜单")
        top_media_series = res_df.groupby('媒体名称')['项目总分'].mean().sort_values(ascending=False).head(10)
        fig_bar = px.bar(
            x=top_media_series.index,
            y=top_media_series.values,
            labels={'x': '媒体名称', 'y': '平均总分'},
            color=top_media_series.values,
            color_continuous_scale='Blues'
        )
        fig_bar.update_layout(showlegend=False, margin=dict(l=20, r=20, t=30, b=40), height=400)
        fig_bar.update_traces(marker_color='#1E88E5')
        st.plotly_chart(fig_bar, use_container_width=True)
        charts['bar'] = fig_bar.to_html(full_html=False, include_plotlyjs='cdn')
        
        st.divider()
        
        # 准备 Top 10 数据用于报告
        df_top_for_report = res_df[['媒体名称', '项目总分', '真需求', '获客效能', '声量']].groupby('媒体名称').mean().sort_values(by='项目总分', ascending=False).head(10).reset_index()
        
        html_report = generate_html_report(
            project_name if project_name else "未命名项目", 
            metrics, 
            charts, 
            df_top_for_report
        )
        
        st.download_button(
            label="📥 下载项目评分报告",
            data=html_report,
            file_name=f"{project_name}_report_view.html" if project_name else "report_view.html",
            mime="text/html",
            type="primary"
        )
